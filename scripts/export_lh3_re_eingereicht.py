#!/usr/bin/env python3
"""Export LH3 Re-Einreichung (Stand: 29.04.2026) nach BARMER-Rückmeldung.

Quelle: lh3/grobkonzept/LH3_Grobkonzept_eingereicht.md (mit eingearbeitetem Feedback F1–F5)
Ziel:   output/LH3_Grobkonzept_Re-Einreichung_2026-04-29.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

YELLOW = "FFD700"
GREEN = "4CAF50"
RED = "C0392B"
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x88, 0x88, 0x88)
ROW_ALT = "FFFBEA"
WHITE = "FFFFFF"


def shade(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_borders(cell, color="DDDDDD", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def style_run(run, *, size=11, bold=False, italic=False, color=None, font="Arial"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc, text="", *, size=11, bold=False, italic=False, color=None, before=3, after=3, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        style_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_runs(doc, parts, *, before=3, after=3):
    """parts: list of dicts {text, bold?, italic?, color?, size?}"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    for part in parts:
        style_run(
            p.add_run(part["text"]),
            size=part.get("size", 11),
            bold=part.get("bold", False),
            italic=part.get("italic", False),
            color=part.get("color"),
        )
    return p


def add_h1(doc, text):
    return add_para(doc, text, size=18, bold=True, color=DARK, before=0, after=6)


def add_h2(doc, text):
    p = add_para(doc, text, size=14, bold=True, color=DARK, before=16, after=7)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), YELLOW)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h3(doc, text):
    return add_para(doc, text, size=12, bold=True, color=DARK, before=10, after=4)


def add_bullet(doc, text, *, prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if prefix:
        style_run(p.add_run(prefix + " "), size=11, bold=True)
    style_run(p.add_run(text), size=11)
    return p


def add_table(doc, headers, rows, widths_cm=None):
    n_cols = len(headers) if headers else len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    if headers:
        hdr = table.add_row()
        for i, h in enumerate(headers):
            c = hdr.cells[i]
            c.text = ""
            shade(c, YELLOW)
            set_borders(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            style_run(p.add_run(h), size=10, bold=True, color=DARK)
            if widths_cm:
                c.width = Cm(widths_cm[i])
    for ri, row in enumerate(rows):
        tr = table.add_row()
        bg = WHITE if ri % 2 == 0 else ROW_ALT
        for i, val in enumerate(row):
            c = tr.cells[i]
            c.text = ""
            shade(c, bg)
            set_borders(c)
            if widths_cm:
                c.width = Cm(widths_cm[i])
            if isinstance(val, list):
                # list of run-dicts
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for part in val:
                    style_run(
                        p.add_run(part["text"]),
                        size=part.get("size", 10),
                        bold=part.get("bold", False),
                        italic=part.get("italic", False),
                        color=part.get("color"),
                    )
            else:
                p = c.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                style_run(p.add_run(str(val)), size=10, color=DARK)
    return table


def add_box(doc, color_hex, label, lines):
    fill = "E8F5E9" if color_hex == GREEN else ("FDEBEA" if color_hex == RED else "FFF9C4")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    shade(cell, fill)
    # left accent border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, sz in (("top", "8"), ("bottom", "8"), ("right", "8"), ("left", "24")):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if label:
        style_run(p.add_run(label + "  "), size=11, bold=True)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        style_run(p.add_run(line), size=11)
    add_para(doc, "", before=0, after=2)


# ── Inhalt ────────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # default style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    # Titel
    add_h1(doc, "DURCHBLICKT! – Lehrkräftehandreichung 3")
    add_para(doc, "„Gesundheitsmythen im Netz – Was stimmt wirklich?“", size=14, bold=True, color=DARK, before=0, after=4)
    add_para(doc, "Re-Einreichung nach BARMER-Rückmeldung  |  Stand: 29.04.2026", size=11, italic=True, color=GRAY, before=0, after=12)

    add_table(doc, [], [
        ["Dokument",     "Grobkonzept Re-Einreichung (nach Phase-4b-Rückmeldung)"],
        ["Autor",        "Michael Kohl"],
        ["Verlag",       "Klett MEX"],
        ["Auftraggeber", "BARMER"],
        ["Zielgruppe",   "Lehrkräfte Sekundarstufe I + II (Grundschul-Erweiterung Phase 6a)"],
        ["Dauer",        "90 Minuten (Doppelstunde)"],
    ], widths_cm=[4.5, 12.5])
    add_para(doc, "")

    # Änderungs-Log
    add_h2(doc, "Änderungs-Log – BARMER-Feedback (Stand: 29.04.2026)")
    add_para(doc, "Eingearbeitet aus den roten Anmerkungen unter Abschnitt 5 „Ablaufstruktur“ in der Genehmigungsvorlage. Quelle: BARMER (Kürzel wrf).", before=2, after=6)
    add_table(doc,
        ["Nr.", "Quelle", "Feedback-Punkt", "Umsetzung", "Status"],
        [
            ["F1", "BARMER (wrf)", "DURCHBLICKT!-QuellenChecker einbinden",
             "Phase 3: Tool-Demo im Plenum + optionale Tool-Anwendung in der Gruppenarbeit", "✓"],
            ["F2", "BARMER (wrf)", "Exkurs „Gesundheitsmythen und KI“",
             "Phase 2, Sek-II-Block: KI-generierte Falschinfos, Deepfakes, ChatGPT ohne Quellen", "✓"],
            ["F3", "BARMER (wrf)", "Liste vertrauenswürdiger Gesundheitsquellen",
             "Phase 3 Mini-Infobox + neuer Abschnitt 7 als Lehrkraft-Referenz", "✓"],
            ["F4", "BARMER (wrf)", "Illusory Truth Effect (Phase 2, Sek II)",
             "In Mechanismen-Liste der Sek-II-Spalte ergänzt, mit Kurzdefinition", "✓"],
            ["F5", "BARMER (wrf)", "Mythen-Labor (von passiv zu aktiv)",
             "Phase 6 Transfer als Sek-II-Variante (Sek I bleibt beim Quiz). Bewusste Umverortung von Phase 2 nach Phase 6.", "✓"],
        ],
        widths_cm=[1.0, 2.4, 4.4, 8.0, 1.2]
    )
    add_para(doc, "")
    add_box(doc, YELLOW, "Hinweis zur Umverortung F5:", [
        "Im DOCX als Phase-2-Anregung notiert; inhaltlich („eigenen viralen Post erstellen“) jedoch eine produktive Transfer-Aktivität.",
        "Die Verortung in Phase 6 wahrt Transfer ≠ Reflexion und vermeidet eine zweite produktive Aktivität in der Hinführung.",
    ])

    # Leitlinien 2026
    add_h2(doc, "Leitlinien 2026 – Erfüllung")
    add_table(doc,
        ["Kriterium", "Umsetzung"],
        [
            ["Binnendifferenzierung", "Durchgehend Sek I / Sek II-Varianten in allen Phasen (Phase 6 zusätzlich Quiz vs. Mythen-Labor)"],
            ["Klarer Fokus",          "Eine Leitfrage: Wie erkenne ich verlässliche Gesundheitsinformationen im Netz – und warum ist das für meine Gesundheit wichtig?"],
            ["Andocken",              "Baut auf LH6, LH19, LH8, LH7, LH21, LH14 auf – ohne inhaltliche Dopplung"],
            ["Gesundheitsbezug",      "Körperliche und mentale Gesundheit in jeder Phase verankert"],
            ["Langlebigkeit",         "Faktencheck-Methoden statt plattformspezifischer Details; Beispiele jederzeit austauschbar"],
            ["Strukturelle Kontinuität", "Bewährter 7-Phasen-Aufbau"],
            ["Einheit",               "90 Minuten"],
        ],
        widths_cm=[5.0, 12.0]
    )
    add_para(doc, "")

    # Arbeitstitel
    add_h2(doc, "Arbeitstitel")
    add_box(doc, YELLOW, "", ["„Gesundheitsmythen im Netz – Was stimmt wirklich?“"])

    # 1. Kernbotschaften
    add_h2(doc, "1. Kernbotschaften")
    add_bullet(doc, "Nicht alles, was viral geht, ist wahr – Gesundheitsmythen verbreiten sich online besonders schnell, weil sie einfache Lösungen versprechen.", prefix="1.")
    add_bullet(doc, "Körperbilder und Gesundheitsideale sind oft konstruiert – Social Media zeigt unrealistische, teilweise gesundheitsschädliche Standards.", prefix="2.")
    add_bullet(doc, "Wissenschaftliche Quellen erkennen schützt – Wer Fakten von Mythen unterscheiden kann, trifft bessere Gesundheitsentscheidungen.", prefix="3.")
    add_bullet(doc, "Kritisches Denken ist Gesundheitsschutz – Medienkompetenz befähigt zu informierten Entscheidungen über den eigenen Körper und schützt vor gefährlichen Trends.", prefix="4.")

    # 2. Anknüpfungspunkte
    add_h2(doc, "2. Anknüpfungspunkte an bestehende Einheiten")
    for line in [
        "LH6 Desinformation im Netz – Methoden zur Faktenprüfung (werden vorausgesetzt / kurz aktiviert)",
        "LH19 Körperbilder im Internet – Unrealistische Körperideale, Bildmanipulation",
        "LH8 Digital stark: Gesundheit entdecken – Ernährung, Bewegung, Schlaf als Grundthemen",
        "LH7 Die Welt der Influencerinnen und Influencer – Werbung vs. Information, kommerzielle Interessen",
        "LH21 Medien und Prävention – Gesundheitsinformationen finden und bewerten",
        "LH14 Challenges im Netz – Gefährliche Trends, Gruppendruck",
    ]:
        add_bullet(doc, line)

    # 3. Gesundheitsbezug
    add_h2(doc, "3. Gesundheitsbezug")
    add_h3(doc, "Körperliche Gesundheit")
    for line in [
        "Gefährliche Gesundheitstrends: extreme Diäten, fragwürdige Fitness-Challenges, „Detox“-Produkte",
        "Risiken durch Falschinformationen zu Ernährung, Nahrungsergänzung, Medikamenten",
        "Unrealistische Körperbilder als Motivation für ungesunde Verhaltensweisen (Essstörungen, Übertraining)",
    ]:
        add_bullet(doc, line)
    add_h3(doc, "Mentale Gesundheit")
    for line in [
        "Druck durch idealisierte Körperbilder und Gesundheitsstandards",
        "Verunsicherung und Stress durch widersprüchliche Gesundheitsinformationen",
        "Selbstwertprobleme durch Vergleich mit „perfekten“ Online-Körpern",
        "Schuldgefühle, wenn man „Gesundheitstrends“ nicht folgt",
    ]:
        add_bullet(doc, line)
    add_h3(doc, "Medienkompetenz als Schutzfaktor")
    for line in [
        "Wissenschaftliche Quellen von Meinungen unterscheiden",
        "Kommerzielle Interessen hinter Gesundheitsbehauptungen erkennen",
        "Selbstbestimmte, informierte Entscheidungen über die eigene Gesundheit treffen",
    ]:
        add_bullet(doc, line)

    # 4. KMK
    add_h2(doc, "4. KMK-Kompetenzbereiche")
    add_runs(doc, [{"text": "Schwerpunkt: ", "bold": True}, {"text": "Bereiche 1, 4 und 6"}])

    # 5. DGK
    add_h2(doc, "5. Dimensionen digitaler Gesundheitskompetenz")
    add_runs(doc, [{"text": "Schwerpunkt: ", "bold": True}, {"text": "Dimensionen 2, 3, 4, 5, 7"}])
    for line in [
        "DGK 2: Verstehen, wie Gesundheitsmythen entstehen und sich verbreiten",
        "DGK 3: Kritisches Bewerten von Gesundheitsinformationen und Quellen",
        "DGK 4: Faktencheck-Methoden anwenden, verlässliche Quellen nutzen",
        "DGK 5: Gesundheit durch informierte Entscheidungen schützen",
        "DGK 7: Eigene wissenschaftsbasierte Quiz-Inhalte erstellen",
    ]:
        add_bullet(doc, line)

    # 6. Ablaufstruktur
    add_h2(doc, "6. Ablaufstruktur (90 Minuten)")
    add_runs(doc, [
        {"text": "Video-Idee: ", "bold": True},
        {"text": "Kurzclip – eine Gesundheitsbehauptung (z. B. „Detox-Tee entgiftet deinen Körper“) geht viral. Gegenschnitt: Eine Ernährungswissenschaftlerin oder ein Ernährungswissenschaftler widerlegt die Aussage in 30 Sekunden. Einstiegsfrage: „Wem glaubst du – und warum?“"},
    ])
    add_para(doc, "")

    NEU_RGB = RGBColor(0xC0, 0x39, 0x2B)

    def neu(text):
        return {"text": " " + text, "size": 9, "italic": True, "color": NEU_RGB}

    phases = [
        # (left_runs, right_runs)
        (
            [
                {"text": "Phase 1: Einstieg – „Glaubst du das?“ (10 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Vorwissen aktivieren, persönliche Relevanz herstellen, Neugier wecken. Bildcollage: 8 Gesundheitsbehauptungen aus Social Media (Mix aus Fakt/Fake) im Plenum – spontane Einschätzung per „Daumen hoch/runter“. Anschließend Think-Pair-Share: „Welche Gesundheitstipps habt ihr online gesehen?“ Brückenfrage zur Überleitung: „Welche dieser Behauptungen könnte eure Gesundheitsentscheidungen beeinflussen?“", "size": 10},
            ],
            [
                {"text": "In ", "size": 10}, {"text": "Sek I", "bold": True, "size": 10},
                {"text": " Fokus auf TikTok, Instagram, YouTube (Fitness, Beauty, „Life Hacks“). In ", "size": 10},
                {"text": "Sek II", "bold": True, "size": 10},
                {"text": " zusätzlich Gesundheitsblogs, Influencer-Marketing, pseudowissenschaftliche Studien.", "size": 10},
            ]
        ),
        (
            [
                {"text": "Phase 2: Hinführung – Warum verbreiten sich Mythen? (15 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Mechanismen der Verbreitung verstehen, Verbindung zu Körperbildern und psychischer Gesundheit herstellen. Kurzvideo/Infografik: „Warum verbreiten sich Gesundheitsmythen so schnell?“ im Plenum. In 2er-Gruppen Gründe sammeln (AB 1): Algorithmen, Emotionen, vereinfachte Botschaften, Körperideale. Sicherung im Plenum: „Welche Rolle spielen Körperbilder?“ – Hinweis: Widersprüchliche Gesundheitsinfos erzeugen Verunsicherung, Stress und Schamgefühle. ", "size": 10},
                {"text": "Grundschule (Kl. 3–4): ", "bold": True, "size": 10},
                {"text": "Statt Kurzvideo/Infografik: Detektiv-Geschichte im Sitzkreis. Lehrkraft erzählt: Lena sieht in einem Video, dass ein bestimmtes Getränk nie krank macht – die Geschichte wandert weiter, niemand prüft sie. Ziel: Stille-Post-Prinzip kindgerecht.", "size": 10},
            ],
            [
                {"text": "In ", "size": 10}, {"text": "Sek I", "bold": True, "size": 10},
                {"text": " konkrete Beispiele (Before-After-Bilder, „Wundermittel“, Challenges). In ", "size": 10},
                {"text": "Sek II", "bold": True, "size": 10},
                {"text": " zusätzlich psychologische Mechanismen: Bestätigungsfehler, Dunning-Kruger-Effekt, Sozialer Beweis, ", "size": 10},
                {"text": "Illusory Truth Effect", "bold": True, "size": 10},
                {"text": " (Wiederholung erzeugt subjektiv Wahrheit – je öfter eine Falschinformation wiederholt wird, desto wahrer erscheint sie)", "size": 10},
                neu("[NEU F4]"),
                {"text": ". ", "size": 10},
                {"text": "Kurz-Exkurs „Gesundheitsmythen und KI“: ", "bold": True, "size": 10},
                {"text": "KI-generierte Bilder/Videos, Deepfakes von Ärztinnen, Ärzten oder Prominenten, ChatGPT-Antworten ohne Quellenangabe und automatisiert produzierte „Health-Hacks“-Clips verstärken Mythen massiv – plausibel klingend, aber ungeprüft", "size": 10},
                neu("[NEU F2]"),
                {"text": ".", "size": 10},
            ]
        ),
        (
            [
                {"text": "Phase 3: Erarbeitung I – Faktencheck-Werkstatt (25 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Praktische Methoden zur Überprüfung von Gesundheitsbehauptungen erlernen und anwenden. QQQQ-Methode (Quelle / Qualität / Querbezüge / Qualitätssiegel) im Plenum anhand von AB 2. ", "size": 10},
                {"text": "Anschließend Demo des DURCHBLICKT!-QuellenCheckers (durch-blickt.de): ", "bold": True, "size": 10},
                {"text": "Gemeinsam wird eine Beispielbehauptung mit dem Tool geprüft – Schritt-für-Schritt-Logik und QQQQ-Bezug parallel sichtbar", "size": 10},
                neu("[NEU F1]"),
                {"text": ". In 4er-Gruppen überprüfen die Lernenden 4 Gesundheitsbehauptungen (AB 3a–d). ", "size": 10},
                {"text": "Optional: ", "bold": True, "size": 10},
                {"text": "SuS prüfen 1–2 ihrer Behauptungen zusätzlich mit dem QuellenChecker und vergleichen Tool-Ergebnis und eigene QQQQ-Einschätzung", "size": 10},
                neu("[NEU F1]"),
                {"text": ". Kurzpräsentation im Plenum (je 3 Min.). ", "size": 10},
                {"text": "Gesundheitsbezug: ", "bold": True, "size": 10},
                {"text": "Unkritisch übernommene Gesundheitstipps können zu schädlichen Entscheidungen führen – die QQQQ-Methode schützt vor gesundheitlichen Fehlentscheidungen. ", "size": 10},
                {"text": "Infobox „Wo finde ich verlässliche Gesundheitsinfos?“: ", "bold": True, "size": 10},
                {"text": "Vertrauenswürdige Anlaufstellen u. a.: gesundheitsinformation.de (IQWiG), rki.de, bzga.de, barmer.de/gesundheit. Vollständige Liste s. Abschnitt 7", "size": 10},
                neu("[NEU F3]"),
                {"text": ". ", "size": 10},
                {"text": "Grundschule (Kl. 3–4): ", "bold": True, "size": 10},
                {"text": "Statt QQQQ-Methode: Detektiv-Arbeitsblatt mit drei kindgerechten Fragen: 1. „Wer sagt das?“ 2. „Warum sagt die Person das?“ 3. „Fragt noch jemand anders?“", "size": 10},
            ],
            [
                {"text": "In ", "size": 10}, {"text": "Sek I", "bold": True, "size": 10},
                {"text": " vereinfachte QQQQ-Checkliste, klare Mythen vs. Fakten, vorgegebene Quellen. In ", "size": 10},
                {"text": "Sek II", "bold": True, "size": 10},
                {"text": " zusätzlich: Studienqualität bewerten, Interessenkonflikte erkennen, eigene Recherche.", "size": 10},
            ]
        ),
        (
            [
                {"text": "Phase 4: Erarbeitung II – Körperbilder & Gesundheitsmythen (15 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Verbindung zwischen konstruierten Körperbildern und Gesundheitsmythen erkennen. In 4er-Gruppen analysieren die Lernenden 3 fiktive Influencer-Profile (Fitness / Ernährung / Beauty) anhand von AB 4. Leitfragen: „Welches Körperbild wird vermittelt? Welche Mythen? Wer profitiert?“", "size": 10},
            ],
            [
                {"text": "In ", "size": 10}, {"text": "Sek I", "bold": True, "size": 10},
                {"text": " Beschreiben und Erkennen von Unrealistischem, persönliche Einschätzung. In ", "size": 10},
                {"text": "Sek II", "bold": True, "size": 10},
                {"text": " zusätzlich: Geschäftsmodelle analysieren, gesellschaftliche Auswirkungen, Bodyshaming, Essstörungen.", "size": 10},
            ]
        ),
        (
            [
                {"text": "Phase 5: Vertiefung – Wer steckt dahinter? (10 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Kommerzielle und algorithmische Interessen hinter Gesundheitsmythen erkennen. Impulsfrage im Plenum: „Wem nützt es, wenn Gesundheitsmythen geglaubt werden?“ Diskussion im Plenum: „Macht es einen Unterschied, ob jemand bezahlt wurde?“ ", "size": 10},
                {"text": "Gesundheitsbezug: ", "bold": True, "size": 10},
                {"text": "Kommerzielle Gesundheitsmythen erzeugen gezielt Unsicherheit über den eigenen Körper – das Erkennen dieser Mechanismen schützt vor Druck, Schamgefühlen und unrealistischen Körperidealen.", "size": 10},
            ],
            [
                {"text": "In ", "size": 10}, {"text": "Sek I", "bold": True, "size": 10},
                {"text": " Werbung erkennen, einfache Interessen benennen. In ", "size": 10},
                {"text": "Sek II", "bold": True, "size": 10},
                {"text": " zusätzlich: EU-Regulierung, Transparenzpflichten, gesellschaftliche Verantwortung von Plattformen.", "size": 10},
            ]
        ),
        (
            [
                {"text": "Phase 6: Transfer (10 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Gelerntes anwenden – von passiver Mythen-Erkennung zu aktivem Mechanismen-Verständnis. ", "size": 10},
                {"text": "Sek I – „Wahr oder Fake?“-Quiz: ", "bold": True, "size": 10},
                {"text": "In 2er-Gruppen erstellen die Lernenden je 2–3 Quiz-Fragen mit Auflösung und kurzer Begründung (AB 5). Die besten Fragen werden im Plenum gespielt. ", "size": 10},
                {"text": "Sek II – Mythen-Labor", "bold": True, "size": 10},
                neu("[NEU F5]"),
                {"text": ": ", "bold": True, "size": 10},
                {"text": "In 2er-Gruppen entwickeln die Lernenden einen eigenen fiktiven viralen Gesundheits-Post (Behauptung + Bild-/Captionidee + Hashtag) inkl. Mythen-Auflösung auf der Rückseite: Welche QQQQ-Schwächen wurden gezielt eingebaut (z. B. fehlende Quelle, emotionale Sprache, scheinbares Qualitätssiegel, Cherry Picking)? Anschließend kurze Galerie-Runde im Plenum (3–4 ausgewählte Posts). ", "size": 10},
                {"text": "Persönliches Faktencheck-Versprechen auf Notizzettel – einheitlich für beide Niveaus.", "size": 10},
            ],
            [
                {"text": "In ", "size": 10}, {"text": "Sek I", "bold": True, "size": 10},
                {"text": " Multiple-Choice-Quiz mit klaren Antworten; Fokus: Wahr oder Fake? In ", "size": 10},
                {"text": "Sek II", "bold": True, "size": 10},
                {"text": " Mythen-Labor: produktive Anwendung der QQQQ-Kriterien durch bewusstes Konstruieren eines glaubwürdig wirkenden Mythos und Reflexion über die eingebauten Manipulationsmechanismen", "size": 10},
                neu("[NEU F5]"),
                {"text": ".", "size": 10},
            ]
        ),
        (
            [
                {"text": "Phase 7: Reflexion (5 Min.) ", "bold": True, "size": 10},
                {"text": "Ziel: ", "bold": True, "size": 10},
                {"text": "Rückblick auf den Lernprozess – kein Produkt, kein Handlungsdruck. Blitzlicht-Abschlussrunde im Plenum: „Was nehme ich heute mit? Was hat mich überrascht? Was beschäftigt mich noch?“ Offene Fragen werden gesammelt – ohne Auflösung. Optional: Kurze anonyme Selbstreflexion: „Welchen Gesundheitstipp werde ich das nächste Mal hinterfragen?“", "size": 10},
            ],
            [
                {"text": "Einheitlich für Sek I und Sek II. Niedrigschwellig, freiwillig, kein Bewertungsdruck.", "size": 10},
            ]
        ),
    ]

    add_table(doc,
        ["Phase / Inhalt", "Sek I / Sek II"],
        [[left, right] for (left, right) in phases],
        widths_cm=[10.0, 7.0]
    )
    add_para(doc, "")
    add_runs(doc, [
        {"text": "Phasenzeiten gesamt: ", "bold": True},
        {"text": "10 + 15 + 25 + 15 + 10 + 10 + 5 = 90 Min."},
    ])

    # 7. Vertrauenswürdige Quellen
    add_h2(doc, "7. Vertrauenswürdige Gesundheitsquellen (Lehrkraft-Referenz)")
    add_runs(doc, [{"text": "[NEU F3]", "italic": True, "color": NEU_RGB, "size": 10}])
    add_para(doc, "Kuratierte Liste valider Anlaufstellen für Faktenchecks und vertiefende Recherche. Sie dient der eigenen Vorbereitung der Lehrkraft und kann den Lernenden in Phase 3 als Aushang/Handout zur Verfügung gestellt werden – ersetzt aber nicht die Anwendung der QQQQ-Methode.")

    add_table(doc,
        ["Quelle", "URL", "Profil"],
        [
            ["IQWiG / gesundheitsinformation.de",                "gesundheitsinformation.de",                  "Unabhängige, evidenzbasierte Gesundheitsinformationen für die Allgemeinbevölkerung"],
            ["Robert Koch-Institut (RKI)",                       "rki.de",                                     "Bundesinstitut für Infektions- und nichtübertragbare Krankheiten"],
            ["Bundeszentrale für gesundheitliche Aufklärung",    "bzga.de",                                    "Prävention, Aufklärung, Jugendgesundheit"],
            ["BARMER Gesundheitsportal",                         "barmer.de/gesundheit",                       "Krankenkassen-Gesundheitsinfos, Themenseiten"],
            ["Stiftung Gesundheitswissen",                       "stiftung-gesundheitswissen.de",              "Patientenorientierte, qualitätsgeprüfte Inhalte"],
            ["MedlinePlus / Cochrane (engl.)",                   "medlineplus.gov / cochrane.org",             "International, systematische Übersichtsarbeiten"],
            ["DURCHBLICKT! QuellenChecker",                      "durch-blickt.de",                            "Interaktives Tool zur Quellenprüfung (in Phase 3 eingebunden)"],
        ],
        widths_cm=[5.5, 4.5, 7.0]
    )
    add_para(doc, "")

    # Footer
    add_para(doc, "DURCHBLICKT! – Digital in eine gesunde Zukunft  |  BARMER / Klett MEX",
             size=9, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=2)
    add_para(doc, "Re-Einreichung LH3  |  Autor: Michael Kohl  |  29.04.2026",
             size=9, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    return doc


if __name__ == "__main__":
    out = Path(__file__).parent.parent / "output" / "LH3_Grobkonzept_Re-Einreichung_2026-04-29.docx"
    out.parent.mkdir(exist_ok=True)
    doc = build()
    doc.save(out)
    print(f"DOCX erstellt: {out}")
